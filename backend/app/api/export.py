from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse, Response
from sqlalchemy.orm import Session
import csv
import json
import io
from docx import Document
from app.core.database import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.chapter import Chapter
from app.models.question import Question, QuestionStatus
from app.models.subject import Subject

router = APIRouter(prefix="/export", tags=["export"])


@router.get("/{chapter_id}")
def export_questions(
    chapter_id: int,
    format: str,  # csv, json, docx
    type: str,    # all, review, almost_forgot, errors
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Verify chapter ownership
    chapter = db.query(Chapter).join(Subject).filter(
        Chapter.id == chapter_id,
        Subject.user_id == current_user.id
    ).first()
    
    if not chapter:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Chapter not found"
        )
        
    # Get base query for questions in chapter
    query = db.query(Question).filter(Question.chapter_id == chapter_id)
    
    # Filter by status type
    if type == "review":
        query = query.filter(Question.status == QuestionStatus.REVIEW)
    elif type == "almost_forgot":
        query = query.filter(Question.status == QuestionStatus.ALMOST_FORGOT)
    elif type == "errors":
        query = query.filter(Question.status == QuestionStatus.ERROR)
    elif type != "all":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid export type: {type}"
        )
        
    questions = query.order_by(Question.question_number.asc()).all()
    filename_base = f"chapter_{chapter_id}_{type}_export"
    
    if format == "json":
        data = [
            {
                "question_number": q.question_number,
                "question_text": q.question_text,
                "option_a": q.option_a,
                "option_b": q.option_b,
                "option_c": q.option_c,
                "option_d": q.option_d,
                "correct_answer": q.correct_answer,
                "status": q.status.value if q.status else "NEW"
            }
            for q in questions
        ]
        json_str = json.dumps(data, indent=2)
        return Response(
            content=json_str,
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename={filename_base}.json"}
        )
        
    elif format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow([
            "Question Number", "Question Text", 
            "Option A", "Option B", "Option C", "Option D", 
            "Correct Answer", "Status"
        ])
        
        for q in questions:
            writer.writerow([
                q.question_number,
                q.question_text,
                q.option_a,
                q.option_b,
                q.option_c,
                q.option_d,
                q.correct_answer,
                q.status.value if q.status else "NEW"
            ])
            
        csv_str = output.getvalue()
        return Response(
            content=csv_str,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename_base}.csv"}
        )
        
    elif format == "docx":
        doc = Document()
        doc.add_heading(f"Exported Questions - {chapter.name}", level=1)
        doc.add_paragraph(f"Export Type: {type.upper()} | Total Questions: {len(questions)}")
        doc.add_paragraph("")
        
        for q in questions:
            p = doc.add_paragraph()
            p.add_run(f"Question {q.question_number}\n").bold = True
            p.add_run(f"{q.question_text}\n")
            p.add_run(f"A. {q.option_a}\n")
            p.add_run(f"B. {q.option_b}\n")
            p.add_run(f"C. {q.option_c}\n")
            p.add_run(f"D. {q.option_d}\n")
            p.add_run(f"Answer: {q.correct_answer.upper()}\n").bold = True
            doc.add_paragraph("") # Spacing
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename_base}.docx"}
        )
        
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid format: {format}"
        )
